"""
File handling and conversion service
"""
import os
import uuid
from typing import Optional, Dict, Any, List
from pathlib import Path
import fitz  # PyMuPDF
import docx
from docx import Document
import nbformat
from nbconvert import HTMLExporter
from PIL import Image
import io
import base64

from ..utils.file_utils import FileHandler, FileValidator

class FileService:
    """Service for file operations and conversions"""
    
    def __init__(self, file_handler: FileHandler):
        self.file_handler = file_handler
    
    def process_uploaded_file(self, file, file_type: str) -> Dict[str, Any]:
        """Process an uploaded file based on its type"""
        try:
            # Save the file
            file_path = self.file_handler.save_file(file, file.filename)
            
            # Process based on file type
            if file_type == 'pdf':
                return self._process_pdf_file(file_path)
            elif file_type == 'python':
                return self._process_python_file(file_path)
            elif file_type == 'jupyter':
                return self._process_jupyter_file(file_path)
            elif file_type == 'document':
                return self._process_document_file(file_path)
            elif file_type == 'image':
                return self._process_image_file(file_path)
            else:
                return {'error': f'Unsupported file type: {file_type}'}
                
        except Exception as e:
            return {'error': f'Error processing file: {str(e)}'}
    
    def _process_pdf_file(self, file_path: str) -> Dict[str, Any]:
        """Process PDF file"""
        try:
            pdf_doc = fitz.open(file_path)
            
            # Get basic information
            info = {
                'filename': os.path.basename(file_path),
                'file_path': file_path,
                'pages': len(pdf_doc),
                'file_size': os.path.getsize(file_path)
            }
            
            # Extract text from first page for preview
            if len(pdf_doc) > 0:
                first_page = pdf_doc[0]
                text = first_page.get_text()
                info['preview_text'] = text[:500] + '...' if len(text) > 500 else text
            
            pdf_doc.close()
            return info
            
        except Exception as e:
            return {'error': f'Error processing PDF: {str(e)}'}
    
    def _process_python_file(self, file_path: str) -> Dict[str, Any]:
        """Process Python file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            return {
                'filename': os.path.basename(file_path),
                'file_path': file_path,
                'content': content,
                'lines': len(content.splitlines()),
                'file_size': os.path.getsize(file_path)
            }
            
        except Exception as e:
            return {'error': f'Error processing Python file: {str(e)}'}
    
    def _process_jupyter_file(self, file_path: str) -> Dict[str, Any]:
        """Process Jupyter notebook file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                notebook_content = f.read()
            
            notebook = nbformat.reads(notebook_content, as_version=4)
            
            # Convert to HTML
            html_exporter = HTMLExporter()
            html_content, _ = html_exporter.from_notebook_node(notebook)
            
            return {
                'filename': os.path.basename(file_path),
                'file_path': file_path,
                'html_content': html_content,
                'cell_count': len(notebook.cells),
                'file_size': os.path.getsize(file_path)
            }
            
        except Exception as e:
            return {'error': f'Error processing Jupyter notebook: {str(e)}'}
    
    def _process_document_file(self, file_path: str) -> Dict[str, Any]:
        """Process Word document file"""
        try:
            doc = Document(file_path)
            
            # Extract text content
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            return {
                'filename': os.path.basename(file_path),
                'file_path': file_path,
                'content': '\n'.join(text_content),
                'paragraphs': len(text_content),
                'file_size': os.path.getsize(file_path)
            }
            
        except Exception as e:
            return {'error': f'Error processing document: {str(e)}'}
    
    def _process_image_file(self, file_path: str) -> Dict[str, Any]:
        """Process image file"""
        try:
            with Image.open(file_path) as img:
                # Convert to base64 for display
                img_buffer = io.BytesIO()
                img.save(img_buffer, format='PNG')
                img_b64 = base64.b64encode(img_buffer.getvalue()).decode()
                
                return {
                    'filename': os.path.basename(file_path),
                    'file_path': file_path,
                    'data': f'data:image/png;base64,{img_b64}',
                    'width': img.width,
                    'height': img.height,
                    'format': img.format,
                    'file_size': os.path.getsize(file_path)
                }
                
        except Exception as e:
            return {'error': f'Error processing image: {str(e)}'}
    
    def convert_pdf_to_word(self, pdf_path: str) -> Optional[str]:
        """Convert PDF to Word document"""
        try:
            pdf_doc = fitz.open(pdf_path)
            doc = Document()
            
            for page_num in range(len(pdf_doc)):
                page = pdf_doc[page_num]
                text = page.get_text()
                
                if page_num > 0:
                    doc.add_page_break()
                
                if text.strip():
                    doc.add_paragraph(text)
            
            # Save converted document
            output_path = self.file_handler.temp_folder / f'converted_{uuid.uuid4()}.docx'
            doc.save(str(output_path))
            pdf_doc.close()
            
            return str(output_path)
            
        except Exception as e:
            print(f"Error converting PDF to Word: {e}")
            return None
    
    def convert_word_to_pdf(self, docx_path: str) -> Optional[str]:
        """Convert Word document to PDF"""
        try:
            doc = Document(docx_path)
            
            # Create new PDF
            pdf_doc = fitz.open()
            page = pdf_doc.new_page()
            
            y_position = 72  # Start 1 inch from top
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    page.insert_text((72, y_position), paragraph.text, fontsize=12)
                    y_position += 20
                    
                    # Add new page if needed
                    if y_position > 700:
                        page = pdf_doc.new_page()
                        y_position = 72
            
            # Save converted PDF
            output_path = self.file_handler.temp_folder / f'converted_{uuid.uuid4()}.pdf'
            pdf_doc.save(str(output_path))
            pdf_doc.close()
            
            return str(output_path)
            
        except Exception as e:
            print(f"Error converting Word to PDF: {e}")
            return None
    
    def convert_images_to_pdf(self, image_paths: List[str]) -> Optional[str]:
        """Convert multiple images to PDF"""
        try:
            pdf_doc = fitz.open()
            
            for image_path in image_paths:
                # Open image
                with Image.open(image_path) as img:
                    # Convert to RGB if necessary
                    if img.mode != 'RGB':
                        img = img.convert('RGB')
                    
                    # Save as temporary file
                    temp_path = self.file_handler.temp_folder / f'temp_{uuid.uuid4()}.jpg'
                    img.save(str(temp_path), 'JPEG')
                    
                    # Add to PDF
                    page = pdf_doc.new_page()
                    page.insert_image(page.rect, filename=str(temp_path))
                    
                    # Clean up temp file
                    os.remove(str(temp_path))
            
            # Save PDF
            output_path = self.file_handler.temp_folder / f'images_to_pdf_{uuid.uuid4()}.pdf'
            pdf_doc.save(str(output_path))
            pdf_doc.close()
            
            return str(output_path)
            
        except Exception as e:
            print(f"Error converting images to PDF: {e}")
            return None
    
    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """Extract all text from PDF"""
        try:
            pdf_doc = fitz.open(pdf_path)
            text_content = []
            
            for page_num in range(len(pdf_doc)):
                page = pdf_doc[page_num]
                text = page.get_text()
                if text.strip():
                    text_content.append(text)
            
            pdf_doc.close()
            return '\n\n'.join(text_content)
            
        except Exception as e:
            print(f"Error extracting text from PDF: {e}")
            return ""
    
    def get_file_preview(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """Get file preview information"""
        try:
            if not os.path.exists(file_path):
                return {'error': 'File not found'}
            
            file_info = self.file_handler.get_file_info(file_path)
            
            preview = {
                'filename': os.path.basename(file_path),
                'file_size': file_info.get('size', 0),
                'mime_type': file_info.get('mime_type', ''),
                'created': file_info.get('created', 0),
                'modified': file_info.get('modified', 0)
            }
            
            # Add type-specific preview
            if file_type == 'pdf':
                preview.update(self._get_pdf_preview(file_path))
            elif file_type == 'image':
                preview.update(self._get_image_preview(file_path))
            elif file_type == 'python':
                preview.update(self._get_python_preview(file_path))
            
            return preview
            
        except Exception as e:
            return {'error': f'Error getting file preview: {str(e)}'}
    
    def _get_pdf_preview(self, file_path: str) -> Dict[str, Any]:
        """Get PDF-specific preview"""
        try:
            pdf_doc = fitz.open(file_path)
            preview = {
                'pages': len(pdf_doc),
                'title': pdf_doc.metadata.get('title', ''),
                'author': pdf_doc.metadata.get('author', '')
            }
            pdf_doc.close()
            return preview
        except Exception:
            return {}
    
    def _get_image_preview(self, file_path: str) -> Dict[str, Any]:
        """Get image-specific preview"""
        try:
            with Image.open(file_path) as img:
                return {
                    'width': img.width,
                    'height': img.height,
                    'format': img.format,
                    'mode': img.mode
                }
        except Exception:
            return {}
    
    def _get_python_preview(self, file_path: str) -> Dict[str, Any]:
        """Get Python file-specific preview"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                lines = f.readlines()
            
            return {
                'lines': len(lines),
                'first_line': lines[0].strip() if lines else '',
                'has_imports': any(line.strip().startswith('import') or line.strip().startswith('from') for line in lines[:10])
            }
        except Exception:
            return {}
    
    def cleanup_temp_files(self) -> int:
        """Clean up temporary files"""
        return self.file_handler.cleanup_temp_files()
