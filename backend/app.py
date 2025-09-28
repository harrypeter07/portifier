#!/usr/bin/env python3
"""
PDF Editor Service - Flask API
A comprehensive REST API for PDF editing, conversion, and analysis
"""

from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import fitz  # PyMuPDF
import json
import os
import io
import base64
import tempfile
from pathlib import Path
import cv2
import numpy as np
from PIL import Image
import re
from dataclasses import dataclass, asdict
from typing import List, Dict, Tuple, Optional
import pandas as pd
import docx
from docx import Document
import pytesseract
from pdf2image import convert_from_path
import nbformat
from nbconvert import HTMLExporter
import zipfile
import shutil

app = Flask(__name__)
CORS(app)

# Configure upload folder
UPLOAD_FOLDER = 'uploads'
TEMP_FOLDER = 'temp'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(TEMP_FOLDER, exist_ok=True)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

@dataclass
class TextElement:
    """Represents a text element with all its formatting properties"""
    text: str
    bbox: Tuple[float, float, float, float]  # x0, y0, x1, y1
    font_name: str
    font_size: float
    font_flags: int
    color: Tuple[int, int, int]
    page_num: int
    block_num: int
    line_num: int
    word_num: int
    element_id: str

class PDFProcessor:
    def __init__(self):
        self.pdf_document = None
        self.text_elements = []
        self.images = []
        self.extracted_colors = set()
        self.extracted_fonts = set()
    
    def load_pdf(self, pdf_path):
        """Load PDF and extract all elements"""
        try:
            self.pdf_document = fitz.open(pdf_path)
            self.extract_text_elements()
            self.extract_images()
            self.analyze_fonts_and_colors()
            return True
        except Exception as e:
            print(f"Error loading PDF: {e}")
            return False
    
    def extract_text_elements(self):
        """Extract all text elements with detailed formatting"""
        self.text_elements = []
        
        for page_num in range(len(self.pdf_document)):
            page = self.pdf_document[page_num]
            blocks = page.get_text("dict")
            
            for block_num, block in enumerate(blocks["blocks"]):
                if "lines" not in block:
                    continue
                
                for line_num, line in enumerate(block["lines"]):
                    for word_num, word in enumerate(line["spans"]):
                        color = word.get("color", 0)
                        rgb_color = (
                            (color >> 16) & 255,
                            (color >> 8) & 255,
                            color & 255
                        )
                        
                        element_id = f"p{page_num}_b{block_num}_l{line_num}_w{word_num}"
                        
                        text_element = TextElement(
                            text=word["text"],
                            bbox=word["bbox"],
                            font_name=word["font"],
                            font_size=word["size"],
                            font_flags=word["flags"],
                            color=rgb_color,
                            page_num=page_num,
                            block_num=block_num,
                            line_num=line_num,
                            word_num=word_num,
                            element_id=element_id
                        )
                        
                        self.text_elements.append(text_element)
    
    def extract_images(self):
        """Extract all images from PDF"""
        self.images = []
        
        for page_num in range(len(self.pdf_document)):
            page = self.pdf_document[page_num]
            image_list = page.get_images()
            
            for img_index, img in enumerate(image_list):
                xref = img[0]
                pix = fitz.Pixmap(self.pdf_document, xref)
                
                if pix.n - pix.alpha < 4:  # GRAY or RGB
                    img_data = pix.tobytes("png")
                    img_b64 = base64.b64encode(img_data).decode()
                    
                    # Get image rectangle
                    rects = page.get_image_rects(xref)
                    bbox = rects[0] if rects else (0, 0, 100, 100)
                    
                    self.images.append({
                        'id': f"img_{page_num}_{img_index}",
                        'page': page_num,
                        'bbox': bbox,
                        'data': img_b64,
                        'xref': xref
                    })
                
                pix = None
    
    def analyze_fonts_and_colors(self):
        """Analyze unique fonts and colors"""
        self.extracted_fonts = set()
        self.extracted_colors = set()
        
        for element in self.text_elements:
            font_info = f"{element.font_name} ({element.font_size}pt)"
            if element.font_flags & 2**4:  # Bold
                font_info += " Bold"
            if element.font_flags & 2**1:  # Italic
                font_info += " Italic"
            
            self.extracted_fonts.add(font_info)
            self.extracted_colors.add(element.color)
    
    def update_text_element(self, element_id, new_text, new_font_size=None, new_color=None):
        """Update a specific text element"""
        try:
            # Find the element
            element = None
            for el in self.text_elements:
                if el.element_id == element_id:
                    element = el
                    break
            
            if not element:
                return False
            
            page = self.pdf_document[element.page_num]
            
            # Remove old text by drawing white rectangle over it
            rect = fitz.Rect(element.bbox)
            page.draw_rect(rect, color=(1, 1, 1), fill=(1, 1, 1))
            
            # Insert new text
            font_size = new_font_size if new_font_size else element.font_size
            color = new_color if new_color else element.color
            color_fitz = (color[0]/255, color[1]/255, color[2]/255)
            
            page.insert_text(
                (element.bbox[0], element.bbox[1] + font_size),
                new_text,
                fontsize=font_size,
                color=color_fitz,
                fontname=element.font_name
            )
            
            # Update element data
            element.text = new_text
            if new_font_size:
                element.font_size = new_font_size
            if new_color:
                element.color = new_color
            
            return True
        except Exception as e:
            print(f"Error updating text: {e}")
            return False
    
    def extract_text_from_images(self):
        """Extract text from images using OCR"""
        ocr_results = []
        
        for img in self.images:
            try:
                # Decode base64 image
                img_data = base64.b64decode(img['data'])
                pil_img = Image.open(io.BytesIO(img_data))
                
                # Convert to numpy array for OCR
                img_array = np.array(pil_img)
                
                # Extract text using Tesseract
                text = pytesseract.image_to_string(img_array)
                
                ocr_results.append({
                    'image_id': img['id'],
                    'page': img['page'],
                    'bbox': img['bbox'],
                    'extracted_text': text.strip()
                })
            except Exception as e:
                print(f"OCR error for image {img['id']}: {e}")
                continue
        
        return ocr_results
    
    def get_page_image(self, page_num, zoom=1.5):
        """Get page as base64 encoded image"""
        try:
            page = self.pdf_document[page_num]
            mat = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat)
            img_data = pix.tobytes("png")
            return base64.b64encode(img_data).decode()
        except Exception as e:
            print(f"Error getting page image: {e}")
            return None

# Global processor instance
pdf_processor = PDFProcessor()

@app.route('/api/upload', methods=['POST'])
def upload_file():
    """Upload and process PDF file"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        if file and file.filename.lower().endswith('.pdf'):
            filename = f"uploaded_{file.filename}"
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            
            # Process PDF
            if pdf_processor.load_pdf(filepath):
                return jsonify({
                    'success': True,
                    'message': 'PDF uploaded and processed successfully',
                    'pages': len(pdf_processor.pdf_document),
                    'text_elements': len(pdf_processor.text_elements),
                    'images': len(pdf_processor.images)
                })
            else:
                return jsonify({'error': 'Failed to process PDF'}), 500
        
        return jsonify({'error': 'Invalid file type'}), 400
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/pdf/info', methods=['GET'])
def get_pdf_info():
    """Get PDF information"""
    try:
        if not pdf_processor.pdf_document:
            return jsonify({'error': 'No PDF loaded'}), 400
        
        return jsonify({
            'pages': len(pdf_processor.pdf_document),
            'fonts': list(pdf_processor.extracted_fonts),
            'colors': [{'rgb': color, 'hex': f"#{color[0]:02x}{color[1]:02x}{color[2]:02x}"} 
                      for color in pdf_processor.extracted_colors],
            'text_elements_count': len(pdf_processor.text_elements),
            'images_count': len(pdf_processor.images)
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/pdf/page/<int:page_num>', methods=['GET'])
def get_page(page_num):
    """Get specific page with all elements"""
    try:
        if not pdf_processor.pdf_document:
            return jsonify({'error': 'No PDF loaded'}), 400
        
        if page_num >= len(pdf_processor.pdf_document):
            return jsonify({'error': 'Page not found'}), 404
        
        # Get page image
        page_image = pdf_processor.get_page_image(page_num)
        
        # Get text elements for this page
        page_elements = [asdict(el) for el in pdf_processor.text_elements 
                        if el.page_num == page_num]
        
        # Get images for this page
        page_images = [img for img in pdf_processor.images if img['page'] == page_num]
        
        return jsonify({
            'page_image': page_image,
            'text_elements': page_elements,
            'images': page_images,
            'page_num': page_num
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/pdf/update-text', methods=['POST'])
def update_text():
    """Update text element"""
    try:
        data = request.json
        element_id = data.get('element_id')
        new_text = data.get('new_text')
        new_font_size = data.get('new_font_size')
        new_color = data.get('new_color')
        
        if not element_id or new_text is None:
            return jsonify({'error': 'Missing required parameters'}), 400
        
        success = pdf_processor.update_text_element(element_id, new_text, new_font_size, new_color)
        
        if success:
            return jsonify({'success': True, 'message': 'Text updated successfully'})
        else:
            return jsonify({'error': 'Failed to update text'}), 500
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/pdf/ocr', methods=['GET'])
def extract_image_text():
    """Extract text from images using OCR"""
    try:
        if not pdf_processor.pdf_document:
            return jsonify({'error': 'No PDF loaded'}), 400
        
        ocr_results = pdf_processor.extract_text_from_images()
        return jsonify({'ocr_results': ocr_results})
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/pdf/save', methods=['GET'])
def save_pdf():
    """Save modified PDF"""
    try:
        if not pdf_processor.pdf_document:
            return jsonify({'error': 'No PDF loaded'}), 400
        
        output_path = os.path.join(TEMP_FOLDER, 'edited_document.pdf')
        pdf_processor.pdf_document.save(output_path)
        
        return send_file(output_path, as_attachment=True, download_name='edited_document.pdf')
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/convert/pdf-to-word', methods=['POST'])
def convert_pdf_to_word():
    """Convert PDF to Word document"""
    try:
        if not pdf_processor.pdf_document:
            return jsonify({'error': 'No PDF loaded'}), 400
        
        doc = Document()
        
        # Extract text and add to Word document
        for page_num in range(len(pdf_processor.pdf_document)):
            page = pdf_processor.pdf_document[page_num]
            text = page.get_text()
            
            if page_num > 0:
                doc.add_page_break()
            
            doc.add_paragraph(text)
        
        output_path = os.path.join(TEMP_FOLDER, 'converted_document.docx')
        doc.save(output_path)
        
        return send_file(output_path, as_attachment=True, download_name='converted_document.docx')
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/convert/word-to-pdf', methods=['POST'])
def convert_word_to_pdf():
    """Convert Word document to PDF"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        if not file.filename.lower().endswith('.docx'):
            return jsonify({'error': 'Invalid file type'}), 400
        
        # Save uploaded Word file
        word_path = os.path.join(TEMP_FOLDER, 'temp_word.docx')
        file.save(word_path)
        
        # Convert to PDF (simplified - you might want to use python-docx2pdf or similar)
        doc = Document(word_path)
        
        # Create new PDF
        pdf_doc = fitz.open()
        page = pdf_doc.new_page()
        
        y_position = 72  # Start 1 inch from top
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                page.insert_text((72, y_position), paragraph.text, fontsize=12)
                y_position += 20
        
        output_path = os.path.join(TEMP_FOLDER, 'converted_from_word.pdf')
        pdf_doc.save(output_path)
        pdf_doc.close()
        
        return send_file(output_path, as_attachment=True, download_name='converted_from_word.pdf')
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/python/view', methods=['POST'])
def view_python_file():
    """View Python file with syntax highlighting"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        if not file.filename.lower().endswith('.py'):
            return jsonify({'error': 'Invalid file type'}), 400
        
        content = file.read().decode('utf-8')
        
        return jsonify({
            'filename': file.filename,
            'content': content,
            'lines': len(content.splitlines())
        })
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/jupyter/view', methods=['POST'])
def view_jupyter_notebook():
    """View Jupyter notebook"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        if not file.filename.lower().endswith('.ipynb'):
            return jsonify({'error': 'Invalid file type'}), 400
        
        # Read notebook content
        notebook_content = file.read().decode('utf-8')
        notebook = nbformat.reads(notebook_content, as_version=4)
        
        # Convert to HTML for display
        html_exporter = HTMLExporter()
        html_content, _ = html_exporter.from_notebook_node(notebook)
        
        return jsonify({
            'filename': file.filename,
            'html_content': html_content,
            'cell_count': len(notebook.cells)
        })
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({'status': 'healthy', 'message': 'PDF Editor API is running'})

@app.route('/api/search-replace', methods=['POST'])
def search_and_replace():
    """Search and replace text across the document"""
    try:
        data = request.json
        search_term = data.get('search_term', '')
        replace_with = data.get('replace_with', '')
        
        if not search_term:
            return jsonify({'error': 'Search term required'}), 400
        
        replacements = 0
        
        for element in pdf_processor.text_elements:
            if search_term in element.text:
                new_text = element.text.replace(search_term, replace_with)
                success = pdf_processor.update_text_element(
                    element.element_id, new_text
                )
                if success:
                    replacements += 1
        
        return jsonify({
            'success': True,
            'replacements': replacements,
            'message': f'Replaced {replacements} occurrences'
        })
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/pdf/add-text', methods=['POST'])
def add_text():
    """Add new text to PDF"""
    try:
        data = request.json
        page_num = data.get('page_num', 0)
        x = data.get('x', 100)
        y = data.get('y', 100)
        text = data.get('text', 'New Text')
        font_size = data.get('font_size', 12)
        color = data.get('color', [0, 0, 0])
        
        page = pdf_processor.pdf_document[page_num]
        color_fitz = (color[0]/255, color[1]/255, color[2]/255)
        
        page.insert_text(
            (x, y),
            text,
            fontsize=font_size,
            color=color_fitz
        )
        
        return jsonify({'success': True, 'message': 'Text added successfully'})
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/pdf/add-image', methods=['POST'])
def add_image():
    """Add image to PDF"""
    try:
        if 'image' not in request.files:
            return jsonify({'error': 'No image provided'}), 400
        
        image_file = request.files['image']
        page_num = int(request.form.get('page_num', 0))
        x = float(request.form.get('x', 100))
        y = float(request.form.get('y', 100))
        width = float(request.form.get('width', 100))
        height = float(request.form.get('height', 100))
        
        # Save image temporarily
        temp_image_path = os.path.join(TEMP_FOLDER, 'temp_image.png')
        image_file.save(temp_image_path)
        
        # Insert image into PDF
        page = pdf_processor.pdf_document[page_num]
        rect = fitz.Rect(x, y, x + width, y + height)
        page.insert_image(rect, filename=temp_image_path)